"""
보고서 검증 Worker

3종 검사:
  1. CF 오류   — DCF 금액 ↔ 현재가치 불일치
  2. 누락 오류 — [표 2] 공란, 투자구조도 누락
  3. NAV 오류  — [표 4] 합계 ↔ [표 5] 투자자산 불일치

결과를 '보고서 오류.xlsx'로 저장
"""

import calendar
import re
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path
from typing import List, Tuple, Dict, Any

import openpyxl
from docx import Document
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from .base import AbstractWorker
from utils.domain import COOP_KW, FX_RE, NUM_RE, DATE_PAREN_RE
from utils.messages import Msg


# ── DOCX 상수 ─────────────────────────────────────────────────

_NS_W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TAG_P   = f"{{{_NS_W}}}p"
_TAG_TBL = f"{{{_NS_W}}}tbl"
_TAG_T   = f"{{{_NS_W}}}t"


# ── DOCX 헬퍼 ─────────────────────────────────────────────────

def _el_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(_TAG_T))

def _load_doc(path):
    doc      = Document(path)
    items    = list(doc.element.body)
    tbl_index = {el: i for i, el in enumerate(e for e in items if e.tag == _TAG_TBL)}
    return doc, items, tbl_index

def _to_int(s: str) -> int:
    v = NUM_RE.sub("", s or "")
    return int(v) if v else 0

def _find_p(items, *needles, start=0):
    for i in range(start, len(items)):
        el = items[i]
        if el.tag == _TAG_P:
            txt = _el_text(el)
            if all(n in txt for n in needles):
                return i, txt
    return None, None


# ── 검사 1: CF 오류 ───────────────────────────────────────────

def _check_cf(doc, items, tbl_index, name: str) -> Dict:
    def _result(status, reason, dcf=0, pv=0, diff=0):
        return {"status": status, "filename": name, "reason": reason,
                "dcf": dcf, "pv": pv, "diff": diff}

    pos4, _ = _find_p(items, "[표 4]", "투자자산 평가내역")
    if pos4 is None:
        return _result("skip", "[표 4] 투자자산 평가내역 없음")

    pos_vii, _ = _find_p(items, "VII", "부록", "대출채권", start=pos4 + 1)
    zone_end = pos_vii if pos_vii is not None else len(items)

    for el in items[pos4:zone_end]:
        if _el_text(el) and "외화환산손익이 반영된 금액" in _el_text(el):
            return _result("skip", "외화환산손익 반영 — 검사 제외")

    tbl4 = None
    for el in items[pos4 + 1:zone_end]:
        if el.tag == _TAG_TBL:
            tbl4 = doc.tables[tbl_index[el]]; break
    if tbl4 is None:
        return _result("skip", "투자자산 표 없음")

    headers  = [c.text.strip() for c in tbl4.rows[0].cells]
    note_col = next((j for j, h in enumerate(headers) if h == "비고"), None)
    amt_col  = next((j for j, h in enumerate(headers) if "평가금액" in h), None)
    if note_col is None or amt_col is None:
        return _result("skip", "비고/평가금액 열 없음")

    if FX_RE.search(headers[amt_col]):
        return _result("skip", f"외화({headers[amt_col]}) — 환율 변환 불가, 검사 제외")

    type_col  = next((j for j, h in enumerate(headers) if "상품유형" in h), None)
    data_rows, is_btl = [], False

    for row in tbl4.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) <= note_col: continue
        if not is_btl and ("BTL" in cells[note_col].upper() or "채권 기대회수액" in cells[note_col]):
            is_btl = True
        data_rows.append(cells)

    dcf_total, has_dcf = 0, False
    for cells in data_rows:
        note_val  = cells[note_col]
        is_common = (type_col is not None and len(cells) > type_col and cells[type_col] == "보통주")
        if note_val == "DCF":
            has_dcf = True
        elif not (is_btl and is_common):
            continue
        if len(cells) > amt_col:
            dcf_total += _to_int(cells[amt_col])

    if not has_dcf:
        return _result("skip", "DCF 종목 없음 — 검사 제외")
    if pos_vii is None:
        return _result("error", "부록 섹션 없음 (현재가치 검증 불가)", dcf_total, 0, dcf_total)

    dm = DATE_PAREN_RE.search(name)
    threshold = None
    if dm:
        ed = date.fromisoformat(dm.group(1))
        y, m = ed.year, ed.month - 3
        if m <= 0: m += 12; y -= 1
        threshold = date(y, m, min(ed.day, calendar.monthrange(y, m)[1]))

    pv_total, old_dates = 0, []
    for el in items[pos_vii + 1:]:
        if el.tag != _TAG_TBL: continue
        t = doc.tables[tbl_index[el]]
        if not t.rows: continue
        hdrs = [c.text.strip() for c in t.rows[0].cells]
        if "현재가치" not in hdrs: continue
        pv_col = hdrs.index("현재가치")
        for row in t.rows[1:]:
            pv_total += _to_int(row.cells[pv_col].text)
            if threshold:
                try:
                    if date.fromisoformat(row.cells[0].text.strip()) < threshold:
                        old_dates.append(row.cells[0].text.strip())
                except ValueError:
                    pass

    diff   = dcf_total - pv_total
    errors = []
    if abs(diff) > 100_000:
        errors.append(f"DCF 불일치 (표4 {dcf_total:,} / 현재가치 {pv_total:,} / 차이 {diff:+,}원)")
    if old_dates:
        extra = f" 외 {len(old_dates) - 1}건" if len(old_dates) > 1 else ""
        errors.append(f"기준일 1분기 이전 일자 포함 ({old_dates[0]}{extra})")

    if errors:
        return _result("error", "[유형 1 오류] " + " / ".join(errors), dcf_total, pv_total, diff)
    return _result("ok", f"이상 없음 (차이 {diff:+,}원)", dcf_total, pv_total, diff)


# ── 검사 2: 누락 오류 ─────────────────────────────────────────

def _check_missing(doc, items, tbl_index, name: str) -> Dict:
    issues, tbl2_found, is_coop = [], False, False

    for i, el in enumerate(items):
        if el.tag == _TAG_P and "[표 2]" in _el_text(el):
            tbl2_found = True
            if any(k in _el_text(el) for k in COOP_KW):
                is_coop = True
            for el2 in items[i + 1:]:
                if el2.tag == _TAG_TBL:
                    t = doc.tables[tbl_index[el2]]
                    if any(k in _el_text(el2) for k in COOP_KW):
                        is_coop = True
                    if not t.rows:
                        issues.append("[표 2] 전체 공란")
                    else:
                        ev = [not row.cells[-1].text.strip()
                              for row in t.rows if len(row.cells) >= 2]
                        if all(ev):   issues.append("[표 2] 전체 공란")
                        elif any(ev): issues.append("[표 2] 일부 값 누락")
                    break
            break
    if not tbl2_found:
        issues.append("[표 2] 없음")

    if not is_coop:
        grm_idx = None
        for i, el in enumerate(items):
            if el.tag == _TAG_P:
                txt = _el_text(el)
                if ("투자구조도" in txt or "투자 구조도" in txt) and "[그림" not in txt:
                    grm_idx = i; break
        if grm_idx is None:
            issues.append("투자구조도 섹션 없음")
        else:
            drawing_tag = f"{{{_NS_W}}}drawing"
            has_img = any(
                items[j].findall(f".//{drawing_tag}")
                for j in range(max(0, grm_idx - 2), min(len(items), grm_idx + 10))
            )
            if not has_img:
                issues.append("투자구조도 이미지 누락")

    status = "error" if issues else "ok"
    reason = ", ".join(issues) if issues else "이상 없음"
    return {"status": status, "filename": name, "reason": reason}


# ── 검사 3: NAV 오류 ─────────────────────────────────────────

def _check_nav(doc, items, tbl_index, name: str) -> Dict:
    def _result(status, reason, t4=0, t5=0, diff=0):
        return {"status": status, "filename": name, "reason": reason,
                "tbl4": t4, "tbl5": t5, "diff": diff}

    def _amt_col(hdrs):
        return next((j for j, h in enumerate(hdrs) if "평가금액" in h and "단위당" not in h), None)

    tbl4_total = None
    for i, el in enumerate(items):
        if el.tag != _TAG_P or "투자자산 평가내역" not in _el_text(el): continue
        for el2 in items[i + 1:]:
            if el2.tag == _TAG_TBL:
                t    = doc.tables[tbl_index[el2]]
                hdrs = [c.text.strip() for c in t.rows[0].cells]
                ac   = _amt_col(hdrs)
                if ac is None: break
                for row in reversed(t.rows):
                    cells = [c.text.strip() for c in row.cells]
                    if cells[0] in ("합계", "합  계", "소계"):
                        tbl4_total = _to_int(cells[ac]); break
                break
        if tbl4_total is not None: break

    if tbl4_total is None:
        return _result("skip", "투자자산 평가내역 합계 없음")

    tbl5_inv = None
    for i, el in enumerate(items):
        if el.tag != _TAG_P or "순자산평가금액" not in _el_text(el): continue
        for el2 in items[i + 1:]:
            if el2.tag == _TAG_TBL:
                t    = doc.tables[tbl_index[el2]]
                hdrs = [c.text.strip() for c in t.rows[0].cells]
                ac   = _amt_col(hdrs)
                if ac is None: break
                for row in t.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if "투자자산" in cells[0] and "유동" not in cells[0] and "기타" not in cells[0]:
                        tbl5_inv = _to_int(cells[ac]); break
                break
        if tbl5_inv is not None: break

    if tbl5_inv is None:
        return _result("skip", "순자산평가금액 표 투자자산 없음", tbl4_total)

    diff = tbl4_total - tbl5_inv
    if abs(diff) <= 5:
        return _result("ok", f"이상 없음 (차이 {diff:+,}원)", tbl4_total, tbl5_inv, diff)
    return _result("error",
                   f"[유형 3 오류] NAV 불일치 (표4 {tbl4_total:,} / 표5 {tbl5_inv:,} / 차이 {diff:+,}원)",
                   tbl4_total, tbl5_inv, diff)


# ── 통합 실행 ─────────────────────────────────────────────────

def run_all_checks(path: Path):
    """Document 1회 로딩으로 3종 검사 실행"""
    name = path.name
    try:
        doc, items, tbl_index = _load_doc(path)
    except Exception as e:
        err = f"처리 오류: {e}"
        return (
            {"status": "error", "filename": name, "reason": err, "dcf": 0, "pv": 0, "diff": 0},
            {"status": "error", "filename": name, "reason": err},
            {"status": "error", "filename": name, "reason": err, "tbl4": 0, "tbl5": 0, "diff": 0},
        )
    return (
        _check_cf(doc, items, tbl_index, name),
        _check_missing(doc, items, tbl_index, name),
        _check_nav(doc, items, tbl_index, name),
    )


# ── 파일명 파싱 ───────────────────────────────────────────────

def _parse_stem(stem: str) -> Dict:
    r = {"업무번호": "", "펀드명": "", "평가기준일": "", "기관명": ""}
    for key, pat in (
        ("업무번호",   r"^(J\d+)_"),
        ("펀드명",     r"\(KISP\)(.*?)\s*평가보고서"),
        ("평가기준일", r"평가보고서\((\d{4}-\d{2}-\d{2})\)"),
        ("기관명",     r"_\[([^\]]+)\]$"),
    ):
        m = re.search(pat, stem)
        if m: r[key] = m.group(1).strip()
    return r


# ── 엑셀 저장 ─────────────────────────────────────────────────

def save_error_excel(folder: Path,
                     cf_rows: List[Dict],
                     miss_rows: List[Dict],
                     nav_rows: List[Dict]) -> Path:
    wb  = openpyxl.Workbook()
    HF  = Font(name="맑은 고딕", bold=True, color="FFFFFF", size=9)
    BF  = Font(name="맑은 고딕", size=9)
    HDR = PatternFill("solid", fgColor="1F3864")
    ERR = PatternFill("solid", fgColor="FCE4D6")
    OK  = PatternFill("solid", fgColor="E2EFDA")
    ALT = PatternFill("solid", fgColor="F5F5F5")
    THIN = Side(style="thin", color="BFBFBF")
    BDR  = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    CTR  = Alignment(horizontal="center", vertical="center")
    LEFT = Alignment(horizontal="left",   vertical="center")
    RGT  = Alignment(horizontal="right",  vertical="center")

    label_map = {"ok": "정상", "skip": "제외", "error": "오류"}
    order_map  = {"error": 0, "ok": 1, "skip": 2}

    # ── 시트 1: 리스트 ──
    ws = wb.active
    ws.title = "리스트"
    ws["B2"] = "당기 보고서 리스트"
    ws["B2"].font = Font(name="맑은 고딕", size=24, bold=True, color="1F3864")
    ws.row_dimensions[2].height = 36

    docx_files = sorted(f for f in folder.glob("*.docx")
                        if not f.stem.endswith("_수정") and not f.name.startswith("~$"))
    cols   = ["B", "C", "D", "E", "F"]
    hdr_v  = ["No.", "업무번호", "펀드명", "평가기준일", "기관명"]
    widths = [6, 14, 55, 14, 18]
    aligns = [CTR, CTR, LEFT, CTR, CTR]

    for h, w, col in zip(hdr_v, widths, cols):
        c = ws[f"{col}4"]
        c.value, c.font, c.fill, c.alignment, c.border = h, HF, HDR, CTR, BDR
        ws.column_dimensions[col].width = w
    ws.row_dimensions[4].height = 16

    for ri, f in enumerate(docx_files, 5):
        p    = _parse_stem(f.stem)
        fill = ALT if (ri - 4) % 2 == 0 else None
        vals = [ri - 4, p["업무번호"], p["펀드명"], p["평가기준일"], p["기관명"]]
        for col, v, al in zip(cols, vals, aligns):
            c = ws[f"{col}{ri}"]
            c.value, c.font, c.alignment, c.border = v, BF, al, BDR
            if fill: c.fill = fill
        ws.row_dimensions[ri].height = 14
    ws.freeze_panes = "B5"
    ws.sheet_view.showGridLines = False

    def _make_sheet(title, headers, widths, aligns, rows, fill_fn):
        s = wb.create_sheet(title=title)
        for ci, (h, w) in enumerate(zip(headers, widths), 1):
            c = s.cell(row=1, column=ci, value=h)
            c.font, c.fill, c.alignment, c.border = HF, HDR, CTR, BDR
            s.column_dimensions[c.column_letter].width = w
        s.row_dimensions[1].height = 16
        for ri, r in enumerate(rows, 2):
            fill    = fill_fn(r)
            num_set = r.get("num_cols", set())
            for ci, (v, al) in enumerate(zip(r["vals"], aligns), 1):
                c = s.cell(row=ri, column=ci, value=v)
                c.font, c.fill, c.alignment, c.border = BF, fill, al, BDR
                if ci in num_set and isinstance(v, int):
                    c.number_format = "#,##0"
            s.row_dimensions[ri].height = 14
        s.freeze_panes = "A2"
        s.sheet_view.showGridLines = False

    def _build(src, k1, k2):
        out = []
        for r in sorted([x for x in src if x["status"] != "skip"], key=lambda x: order_map[x["status"]]):
            out.append({
                "vals": [r["filename"], label_map[r["status"]],
                         r.get(k1) or "", r.get(k2) or "",
                         r.get("diff") or "", r["reason"]],
                "status":   r["status"],
                "num_cols": {3, 4, 5},
            })
        return out

    _make_sheet("1. CF 오류",
                ["파일명", "결과", "표4 DCF 합계(원)", "현재가치 합계(원)", "차이(원)", "내용"],
                [65, 10, 22, 22, 18, 70], [LEFT, CTR, RGT, RGT, RGT, LEFT],
                _build(cf_rows, "dcf", "pv"),
                lambda r: ERR if r["status"] == "error" else OK)

    miss_data = [{"vals": [r["filename"], "오류", r["reason"]], "status": "error"}
                 for r in miss_rows if r["status"] == "error"]
    _make_sheet("2. 누락 오류",
                ["파일명", "결과", "누락 항목"],
                [65, 10, 60], [LEFT, CTR, LEFT],
                miss_data, lambda r: ERR)

    _make_sheet("3. NAV 오류",
                ["파일명", "결과", "표4 합계(원)", "표5 투자자산(원)", "차이(원)", "내용"],
                [65, 10, 22, 22, 14, 60], [LEFT, CTR, RGT, RGT, RGT, LEFT],
                _build(nav_rows, "tbl4", "tbl5"),
                lambda r: ERR if r["status"] == "error" else OK)

    out = folder / "보고서 오류.xlsx"
    wb.save(str(out))
    return out


# ── Worker ────────────────────────────────────────────────────

class ReportValidatorWorker(AbstractWorker):
    """보고서 폴더 안 .docx 파일에 3종 검사를 실행하고 엑셀로 저장"""

    def __init__(self, folder: Path):
        super().__init__()
        self.folder = folder
        self.files: List[Path] = []

    def validate_inputs(self) -> bool:
        if not self.folder or not self.folder.is_dir():
            self.finished_error.emit(f"유효하지 않은 폴더: {self.folder}")
            return False
        return True

    def prepare(self) -> bool:
        self.files = sorted(
            f for f in self.folder.glob("*.docx")
            if not f.stem.endswith("_수정") and not f.name.startswith("~$")
        )
        if not self.files:
            self.finished_error.emit("선택 폴더에 검사할 .docx 파일이 없습니다.")
            return False
        self.stats["total"] = len(self.files)
        self.emit_message(f"📁 {len(self.files)}개 파일 발견")
        return True

    def execute(self):
        cf_rows, miss_rows, nav_rows = [], [], []
        cf_counts = {"ok": 0, "skip": 0, "error": 0}
        miss_err = nav_err = 0

        for i, f in enumerate(self.files, 1):
            if not self.check_running():
                return

            self.emit_message(f"🔍 {f.name}", "info")

            cf, miss, nav = run_all_checks(f)

            cf_counts[cf["status"]] += 1
            cf_rows.append(cf)
            if cf["status"] == "ok":
                self.emit_message(f"  ✔ CF 정상 (차이 {cf['diff']:+,}원)", "success")
            elif cf["status"] == "skip":
                self.emit_message(f"  — CF {cf['reason']}", "warning")
            else:
                self.emit_message(f"  ✖ CF {cf['reason']}", "error")

            miss_rows.append(miss)
            if miss["status"] == "error":
                self.emit_message(f"  ✖ 누락 {miss['reason']}", "error")
                miss_err += 1

            nav_rows.append(nav)
            if nav["status"] == "error":
                self.emit_message(f"  ✖ NAV {nav['reason']}", "error")
                nav_err += 1

            self.emit_progress(i, self.stats["total"])

        # 결과를 finalize에서 사용할 수 있도록 인스턴스에 저장
        self._result = {
            "cf_rows":   cf_rows,
            "miss_rows": miss_rows,
            "nav_rows":  nav_rows,
            "cf_counts": cf_counts,
            "miss_err":  miss_err,
            "nav_err":   nav_err,
        }

    def finalize(self):
        """execute 완료 후 run()이 호출 — 시그널은 여기서만 emit"""
        r = getattr(self, "_result", None)
        if r is None:
            return  # check_running() 으로 중단된 경우

        self.emit_message("─" * 50, "info")
        self.emit_message(
            f"CF — 정상 {r['cf_counts']['ok']} / 제외 {r['cf_counts']['skip']} / 오류 {r['cf_counts']['error']}", "info"
        )
        self.emit_message(f"누락 오류 {r['miss_err']}건 | NAV 오류 {r['nav_err']}건", "info")

        try:
            out = save_error_excel(self.folder, r["cf_rows"], r["miss_rows"], r["nav_rows"])
            self.emit_message(f"📄 결과 저장 → {out.name}", "success")
            self.finished_success.emit(
                f"검사 완료 — CF오류 {r['cf_counts']['error']}건 / "
                f"누락 {r['miss_err']}건 / NAV {r['nav_err']}건\n"
                f"결과 파일: {out.name}"
            )
        except Exception as e:
            self.emit_message(f"❌ 엑셀 저장 실패: {e}", "error")
            self.finished_error.emit(f"엑셀 저장 실패: {e}")
